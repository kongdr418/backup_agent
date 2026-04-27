"""
习题集生成器
生成选择/填空/简答习题集
支持 Markdown 和 Word (.docx) 格式
"""

import os
import re
from datetime import datetime
from typing import Generator, Optional


class ExerciseGenerator:
    """习题集生成器"""

    def __init__(self, output_dir: str = "generated_exercises"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def parse_exercise_request(self, message: str) -> Optional[dict]:
        """
        解析习题集生成请求

        支持的格式：
        - 习题集：Python基础
        - 生成习题：机器学习
        - 习题：深度学习章节2
        - 习题集：Python基础 docx
        - 习题集docx：Python基础
        """
        # 检测是否指定 docx 格式
        output_format = "md"
        if re.search(r'习题集.*docx|docx.*习题集', message, re.IGNORECASE):
            output_format = "docx"
            message = re.sub(r'\s*docx\s*', '', message, flags=re.IGNORECASE)

        patterns = [
            r'习题集[：:]\s*(.+)',
            r'生成习题[：:]\s*(.+)',
            r'习题[：:]\s*(.+)',
            r'练习题[：:]\s*(.+)',
            r'生成练习[：:]\s*(.+)',
        ]

        for pattern in patterns:
            match = re.search(pattern, message, re.IGNORECASE)
            if match:
                return {
                    'topic': match.group(1).strip(),
                    'type': 'exercise',
                    'format': output_format
                }
        return None

    def generate_exercise_stream(self, topic: str, output_format: str = "md") -> Generator[dict, None, None]:
        """流式生成习题集

        Args:
            topic: 主题
            output_format: 输出格式，"md" 或 "docx"，默认为 "md"
        """
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            result_data = {'topic': topic, 'filepath': '', 'format': output_format}

            yield {
                'step': 'start',
                'progress': 0,
                'status': 'running',
                'data': {},
                'message': f'🚀 开始生成习题集：{topic}'
            }

            content = self._generate_exercise_content(topic)
            result_data['content'] = content

            yield {
                'step': 'generate_content',
                'progress': 60,
                'status': 'running',
                'data': {},
                'message': '✅ 习题生成完成，正在保存...'
            }

            safe_topic = re.sub(r'[^\w一-鿿]+', '_', topic)[:50]

            if output_format == "docx":
                filename = f"习题集_{safe_topic}_{timestamp}.docx"
                filepath = os.path.join(self.output_dir, filename)
                self._save_as_docx(content, filepath, topic)
            else:
                filename = f"习题集_{safe_topic}_{timestamp}.md"
                filepath = os.path.join(self.output_dir, filename)
                header = f"""---
title: {topic} - 习题集
created: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
type: exercise
---

"""
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(header + content)

            result_data['filepath'] = filepath

            yield {
                'step': 'complete',
                'progress': 100,
                'status': 'completed',
                'data': result_data,
                'message': f'🎉 习题集生成完成！格式：{output_format.upper()}'
            }

        except Exception as e:
            yield {
                'step': 'error',
                'progress': 0,
                'status': 'error',
                'data': {'error': str(e)},
                'message': f'❌ 生成失败: {e}'
            }

    def _generate_exercise_content(self, topic: str) -> str:
        """调用 AI 生成习题内容"""
        prompt = f"""请为"{topic}"生成一份完整的习题集。

要求：
1. 包含三种题型：选择题、填空题、简答题
2. 每种题型至少5-10道题
3. 题目难度适中，覆盖核心知识点
4. 包含答案（单独放在最后）
5. 使用 Markdown 格式输出

请按以下格式返回：

# {topic} - 习题集

## 一、选择题（每题X分，共X道）

**1. （难度：易/中/难）**
A. 选项1
B. 选项2
C. 选项3
D. 选项4
【答案：A】

---

**2. ...**

---

## 二、填空题（每题X分，共X道）

**1. （难度：易/中/难）**
【答案：XXX】

---

**2. ...**

---

## 三、简答题（每题X分，共X道）

**1. （难度：易/中/难）**
【参考答案】
- 要点1
- 要点2
- 要点3

---

**2. ...**

---

## 答案汇总

### 选择题答案
1. A
2. B
3. C
...

### 填空题答案
1. XXX
2. XXX
...

请直接返回 Markdown 内容，不要其他说明。"""

        return self._call_llm(prompt)

    def _save_as_docx(self, content: str, filepath: str, topic: str):
        """直接构建 docx 格式的习题集文档"""
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # 设置默认字体（微软雅黑，支持中文）
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # 页面边距
        section = doc.sections[0]
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        # 解析 Markdown 内容
        lines = content.split('\n')
        current_section = None  # 当前题型：choice, fill, essay, answer
        choice_questions = []   # 选择题列表
        fill_questions = []    # 填空题列表
        essay_questions = []    # 简答题列表
        answers = {}           # 答案汇总
        current_question = None
        current_answer = None

        # 解析内容
        for line in lines:
            line = line.strip()
            if not line or line.startswith('```'):
                continue

            # 标题处理
            if line.startswith('# '):
                # 主标题 - 添加文档标题
                p = doc.add_heading(line[2:], level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(18)
            elif line.startswith('## 一、选择题') or line.startswith('## 一，选择题'):
                current_section = 'choice'
            elif line.startswith('## 二、填空题') or line.startswith('## 二，填空题'):
                current_section = 'fill'
            elif line.startswith('## 三、简答题') or line.startswith('## 三，简答题'):
                current_section = 'essay'
            elif line.startswith('## 答案汇总') or line.startswith('## 答案'):
                current_section = 'answer'
            elif line.startswith('### 选择题答案') or line.startswith('### 选择题'):
                current_section = 'answer_choice'
            elif line.startswith('### 填空题答案') or line.startswith('### 填空题'):
                current_section = 'answer_fill'
            # 题目处理
            elif line.startswith('**') and '**' in line[2:]:
                # 提取题目编号
                match = re.match(r'\*\*(\d+)\.', line)
                if match:
                    q_num = match.group(1)
                    q_text = re.sub(r'\*\*\d+\.\s*', '', line).strip()
                    q_text = q_text.replace('**', '')
                    current_question = {'num': q_num, 'text': q_text, 'options': [], 'answer': ''}
                    if current_section == 'choice':
                        choice_questions.append(current_question)
                    elif current_section == 'fill':
                        fill_questions.append(current_question)
                    elif current_section == 'essay':
                        essay_questions.append(current_question)
            # 选项处理
            elif re.match(r'^[A-D][\.．]\s*', line):
                if current_question and current_section == 'choice':
                    option_text = re.sub(r'^[A-D][\.．]\s*', '', line)
                    current_question['options'].append(line)
            # 答案处理
            elif line.startswith('【答案') or line.startswith('【参考答案'):
                if current_question:
                    answer_match = re.search(r'[【\[](答案|参考答案)[：:\]]\s*(.+)', line)
                    if answer_match:
                        current_question['answer'] = answer_match.group(2).strip()
            # 要点列表（简答题答案要点）
            elif line.startswith('- ') and current_section == 'essay' and current_question:
                if 'points' not in current_question:
                    current_question['points'] = []
                current_question['points'].append(line[2:].strip())

        # ========== 构建文档 ==========

        # 添加习题集标题
        title = doc.add_heading(f'{topic} - 习题集', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加元信息
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = meta.add_run(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        run.font.size = Pt(9)
        run.font.color.rgb = None

        # ========== 一、选择题 ==========
        if choice_questions:
            doc.add_heading('一、选择题', level=1)
            for q in choice_questions:
                # 题目
                p = doc.add_paragraph()
                run = p.add_run(f"{q['num']}. {q['text']}")
                run.bold = True
                run.font.size = Pt(11)

                # 选项
                for opt in q.get('options', []):
                    doc.add_paragraph(opt, style='List Bullet')
                # 答案
                if q.get('answer'):
                    p = doc.add_paragraph()
                    run = p.add_run(f"【答案：{q['answer']}】")
                    run.font.color.rgb = None
                doc.add_paragraph()  # 空行

        # ========== 二、填空题 ==========
        if fill_questions:
            doc.add_heading('二、填空题', level=1)
            for q in fill_questions:
                p = doc.add_paragraph()
                run = p.add_run(f"{q['num']}. {q['text']}")
                run.bold = True

                if q.get('answer'):
                    p = doc.add_paragraph()
                    run = p.add_run(f"【答案：{q['answer']}】")
                doc.add_paragraph()

        # ========== 三、简答题 ==========
        if essay_questions:
            doc.add_heading('三、简答题', level=1)
            for q in essay_questions:
                p = doc.add_paragraph()
                run = p.add_run(f"{q['num']}. {q['text']}")
                run.bold = True

                # 要点列表
                if q.get('points'):
                    for point in q['points']:
                        doc.add_paragraph(point, style='List Bullet')

                if q.get('answer'):
                    p = doc.add_paragraph()
                    run = p.add_run(f"【参考答案】")
                    run.bold = True
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.5)
                    run = p.add_run(q['answer'])
                doc.add_paragraph()

        # ========== 答案汇总 ==========
        doc.add_heading('答案汇总', level=1)

        # 选择题答案表
        if choice_questions:
            doc.add_heading('选择题答案', level=2)
            table = doc.add_table(rows=len(choice_questions)+1, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # 表头
            hdr = table.rows[0].cells
            hdr[0].text = '题号'
            hdr[1].text = '答案'
            for cell in hdr:
                cell.paragraphs[0].runs[0].bold = True

            # 数据行
            for i, q in enumerate(choice_questions):
                row = table.rows[i+1].cells
                row[0].text = q['num']
                row[1].text = q['answer']

        # 填空题答案表
        if fill_questions:
            doc.add_heading('填空题答案', level=2)
            table = doc.add_table(rows=len(fill_questions)+1, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            hdr = table.rows[0].cells
            hdr[0].text = '题号'
            hdr[1].text = '答案'
            for cell in hdr:
                cell.paragraphs[0].runs[0].bold = True

            for i, q in enumerate(fill_questions):
                row = table.rows[i+1].cells
                row[0].text = q['num']
                row[1].text = q['answer']

        doc.save(filepath)

    def _call_llm(self, prompt: str) -> str:
        """调用大语言模型"""
        from openai import OpenAI
        client = OpenAI(
            api_key=os.environ.get('DEEPSEEK_API_KEY', ''),
            base_url="https://api.deepseek.com"
        )
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": "你是一位专业的教育工作者，擅长编写高质量的习题集。"},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=4000
        )
        return response.choices[0].message.content


if __name__ == "__main__":
    generator = ExerciseGenerator()
    for update in generator.generate_exercise_stream("Python基础语法"):
        print(f"[{update.get('progress', 0)}%] {update['message']}")
